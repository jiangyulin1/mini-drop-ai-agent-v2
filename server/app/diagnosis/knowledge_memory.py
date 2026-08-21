"""Operator knowledge and durable per-Case retrospective memory."""

from __future__ import annotations

import hashlib
import io
import re
import secrets
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from sqlalchemy import func, or_

from server.app.database import new_session
from server.app.diagnosis.embeddings import embed_documents, embed_query, embedding_status
from server.app.models import CaseMemoryModel, KnowledgeChunkModel, KnowledgeDocumentModel
from server.app.state_machine import now_utc


MAX_DOCUMENT_BYTES = 5 * 1024 * 1024
MAX_CONTENT_CHARS = 200_000
CHUNK_TARGET_CHARS = 1_000
CHUNK_OVERLAP_CHARS = 160
TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".json", ".yaml", ".yml", ".csv", ".log", ".xml"}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | {".docx"}


def _vector_dot(left: list[float], right: list[float]) -> float:
    return sum(float(a) * float(b) for a, b in zip(left, right))


def extract_document_text(filename: str, content: bytes) -> str:
    if len(content) > MAX_DOCUMENT_BYTES:
        raise ValueError("DOCUMENT_TOO_LARGE")
    suffix = Path(filename or "").suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError("UNSUPPORTED_DOCUMENT_TYPE")
    if suffix == ".docx":
        try:
            document = Document(io.BytesIO(content))
            blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            blocks.extend(
                " | ".join(cell.text.strip() for cell in row.cells)
                for table in document.tables for row in table.rows
                if any(cell.text.strip() for cell in row.cells)
            )
            value = "\n".join(blocks)
        except (KeyError, ValueError, OSError) as error:
            raise ValueError("INVALID_DOCX") from error
    else:
        value = _decode_text(content)
    value = value.replace("\x00", "").strip()
    if not value:
        raise ValueError("DOCUMENT_HAS_NO_TEXT")
    return value[:MAX_CONTENT_CHARS]


def create_knowledge_document(
    *, tenant_id: str, case_id: str | None, scope: str, title: str,
    filename: str | None, media_type: str, content_text: str, created_by: str,
    kind: str = "DOCUMENT",
) -> dict[str, Any]:
    normalized_scope = str(scope or "CASE").upper()
    if normalized_scope not in {"CASE", "GLOBAL"}:
        raise ValueError("INVALID_KNOWLEDGE_SCOPE")
    if normalized_scope == "CASE" and not case_id:
        raise ValueError("CASE_ID_REQUIRED")
    normalized = str(content_text or "").strip()[:MAX_CONTENT_CHARS]
    if not normalized:
        raise ValueError("KNOWLEDGE_CONTENT_REQUIRED")
    digest = hashlib.sha256(normalized.encode("utf-8")).hexdigest()
    now = now_utc()
    with new_session() as session:
        query = session.query(KnowledgeDocumentModel).filter(
            KnowledgeDocumentModel.tenant_id == tenant_id,
            KnowledgeDocumentModel.content_sha256 == digest,
        )
        query = query.filter(
            KnowledgeDocumentModel.case_id == case_id
            if case_id else KnowledgeDocumentModel.case_id.is_(None)
        )
        existing = query.first()
        if existing is not None:
            existing.status = "ACTIVE"
            existing.updated_at = now
            if not session.query(KnowledgeChunkModel).filter(
                KnowledgeChunkModel.document_id == existing.document_id,
            ).first():
                _store_chunks(session, existing, now)
            session.commit()
            return _document_dict(session, existing, include_content=True)
        row = KnowledgeDocumentModel(
            document_id=f"knowledge-{secrets.token_hex(12)}",
            tenant_id=tenant_id,
            case_id=case_id if normalized_scope == "CASE" else None,
            scope=normalized_scope,
            kind=str(kind or "DOCUMENT").upper(),
            title=(str(title or filename or "未命名知识").strip() or "未命名知识")[:256],
            filename=(str(filename).strip()[:512] if filename else None),
            media_type=(str(media_type or "text/plain")[:128]),
            content_text=normalized,
            content_sha256=digest,
            status="ACTIVE",
            created_by=str(created_by or "local-development")[:128],
            created_at=now,
            updated_at=now,
        )
        session.add(row)
        session.flush()
        _store_chunks(session, row, now)
        session.commit()
        return _document_dict(session, row, include_content=True)


def list_knowledge_documents(tenant_id: str, case_id: str | None = None) -> list[dict[str, Any]]:
    with new_session() as session:
        query = session.query(KnowledgeDocumentModel).filter(
            KnowledgeDocumentModel.tenant_id == tenant_id,
        )
        if case_id:
            query = query.filter(or_(
                KnowledgeDocumentModel.scope == "GLOBAL",
                KnowledgeDocumentModel.case_id == case_id,
            ))
        else:
            query = query.filter(KnowledgeDocumentModel.scope == "GLOBAL")
        rows = query.order_by(KnowledgeDocumentModel.updated_at.desc()).limit(200).all()
        return [_document_dict(session, row) for row in rows]


def update_knowledge_document(
    document_id: str, tenant_id: str, *, status: str | None = None, title: str | None = None,
) -> dict[str, Any] | None:
    with new_session() as session:
        row = session.query(KnowledgeDocumentModel).filter(
            KnowledgeDocumentModel.document_id == document_id,
            KnowledgeDocumentModel.tenant_id == tenant_id,
        ).first()
        if row is None:
            return None
        if status is not None:
            normalized = status.upper()
            if normalized not in {"ACTIVE", "ARCHIVED"}:
                raise ValueError("INVALID_KNOWLEDGE_STATUS")
            row.status = normalized
        if title is not None and title.strip():
            row.title = title.strip()[:256]
        row.updated_at = now_utc()
        session.commit()
        return _document_dict(session, row)


def retrieve_user_knowledge(
    query: str, *, tenant_id: str, case_id: str | None, limit: int = 5,
) -> list[dict[str, Any]]:
    """Retrieve ranked chunks on demand; document bodies never enter CaseContext."""
    with new_session() as session:
        base = session.query(KnowledgeChunkModel, KnowledgeDocumentModel).join(
            KnowledgeDocumentModel,
            KnowledgeDocumentModel.document_id == KnowledgeChunkModel.document_id,
        ).filter(
            KnowledgeDocumentModel.tenant_id == tenant_id,
            KnowledgeDocumentModel.status == "ACTIVE",
            or_(
                KnowledgeDocumentModel.scope == "GLOBAL",
                KnowledgeDocumentModel.case_id == case_id,
            ),
        )
        if not str(query or "").strip():
            return []
        if embedding_status()["provider"] == "lexical":
            # Keep the safe default fully local and cross-dialect: no embedding
            # call, no pgvector expression, and a hard cap on materialized rows.
            candidates = base.order_by(KnowledgeChunkModel.chunk_id).limit(500).all()
            query_terms = Counter(_tokens(query))

            def lexical_score(
                pair: tuple[KnowledgeChunkModel, KnowledgeDocumentModel],
            ) -> int:
                frequencies = pair[0].term_frequencies or {}
                return sum(
                    min(count, max(0, int(frequencies.get(term) or 0)))
                    for term, count in query_terms.items()
                )

            ranked = sorted(
                candidates,
                key=lambda pair: (-lexical_score(pair), pair[0].chunk_id),
            )
            return [
                _chunk_result(float(lexical_score(pair)), pair[0], pair[1])
                for pair in ranked[:_limit(limit)]
            ]
        query_vector = embed_query(query)
        candidate_limit = max(20, _limit(limit) * 8)
        if session.bind.dialect.name == "postgresql":
            vector_distance = KnowledgeChunkModel.embedding.cosine_distance(query_vector)
            vector_rows = base.order_by(vector_distance).limit(candidate_limit).all()
        else:
            candidates = base.limit(500).all()
            vector_rows = sorted(
                candidates,
                key=lambda pair: -_vector_dot(query_vector, pair[0].embedding or []),
            )[:candidate_limit]

        lexical_rows: list[tuple[KnowledgeChunkModel, KnowledgeDocumentModel]] = []
        if session.bind.dialect.name == "postgresql":
            lexical_query = " ".join(dict.fromkeys(_tokens(query)))
            if lexical_query:
                ts_query = func.plainto_tsquery("simple", lexical_query)
                ts_document = func.to_tsvector("simple", KnowledgeChunkModel.lexical_text)
                lexical_rows = base.filter(ts_document.op("@@")(ts_query)).order_by(
                    func.ts_rank_cd(ts_document, ts_query).desc(),
                ).limit(candidate_limit).all()
        else:
            query_terms = set(_tokens(query))
            lexical_rows = sorted(
                vector_rows,
                key=lambda pair: -len(query_terms & set((pair[0].term_frequencies or {}).keys())),
            )

        # Reciprocal-rank fusion combines mature vector and PostgreSQL FTS
        # retrievers without assuming their raw score scales are comparable.
        fused: dict[str, dict[str, Any]] = {}
        for source, weight in ((vector_rows, 0.7), (lexical_rows, 0.3)):
            for rank, (chunk, document) in enumerate(source, start=1):
                entry = fused.setdefault(chunk.chunk_id, {
                    "score": 0.0, "chunk": chunk, "document": document,
                })
                entry["score"] += weight / (60 + rank)
        ranked = sorted(fused.values(), key=lambda item: (-item["score"], item["chunk"].chunk_id))
        return [
            _chunk_result(item["score"], item["chunk"], item["document"])
            for item in ranked[:_limit(limit)]
        ]


def retrieve_case_memory(query: str, *, tenant_id: str, case_id: str | None, limit: int = 3) -> list[dict[str, Any]]:
    if not case_id:
        return []
    with new_session() as session:
        row = session.query(CaseMemoryModel).filter(
            CaseMemoryModel.case_id == case_id,
            CaseMemoryModel.tenant_id == tenant_id,
        ).first()
        if row is None or not row.summary_text.strip():
            return []
        query_terms = set(_tokens(query))
        ranked: list[tuple[int, int, int, str]] = []
        for index, (start, end, content) in enumerate(chunk_text(row.summary_text)):
            score = len(query_terms & set(_tokens(content)))
            if score:
                ranked.append((score, start, end, content))
        ranked.sort(key=lambda item: (-item[0], item[1]))
        return [{
            "knowledge_id": row.memory_id,
            "document_id": row.memory_id,
            "chunk_id": f"{row.memory_id}:chunk:{start}",
            "chunk_index": index,
            "title": "当前会话复盘记忆",
            "content": content,
            "start_offset": start,
            "end_offset": end,
            "score": float(score),
            "source": "case_memory",
            "scope": "CASE",
            "evidence_refs": row.evidence_refs or [],
            "caveat": "复盘记忆用于召回历史过程，不能替代当前运行时 Evidence。",
        } for index, (score, start, end, content) in enumerate(ranked[:_limit(limit)])]


def get_case_memory(repo: Any, case_id: str, tenant_id: str, *, refresh_auto: bool = False) -> dict[str, Any]:
    with new_session() as session:
        row = session.query(CaseMemoryModel).filter(
            CaseMemoryModel.case_id == case_id,
            CaseMemoryModel.tenant_id == tenant_id,
        ).first()
        if row is not None and not (refresh_auto and row.auto_capture):
            return row.to_dict()
    return refresh_case_memory(repo, case_id, tenant_id)


def refresh_case_memory(repo: Any, case_id: str, tenant_id: str) -> dict[str, Any]:
    case = repo.get_incident_case(case_id, tenant_id)
    if case is None:
        raise ValueError("CASE_NOT_FOUND")
    events = repo.list_case_events(case_id, tenant_id, limit=500) or []
    messages = repo.list_assistant_messages(case_id, tenant_id) if hasattr(repo, "list_assistant_messages") else []
    evidence = repo.list_case_evidence(case_id, tenant_id, limit=200)
    graph = repo.get_case_hypothesis_graph(case_id, tenant_id) or {}
    conclusion = repo.get_conclusion(case_id, tenant_id) if hasattr(repo, "get_conclusion") else None
    recommendations = repo.list_repair_recommendations(case_id, tenant_id) if hasattr(
        repo, "list_repair_recommendations",
    ) else []
    active_hypotheses = [
        item for item in graph.get("hypotheses") or []
        if str(item.get("status") or "") not in {"RULED_OUT", "REJECTED"}
    ]
    evidence_refs = list(dict.fromkeys(
        str(item.get("evidence_id") or "") for item in evidence if item.get("evidence_id")
    ))
    latest_message = str(messages[-1].get("content") or "").strip() if messages else ""
    conclusion_text = str(
        (conclusion or {}).get("report_text")
        or (conclusion or {}).get("abstention_reason")
        or latest_message
        or "尚未形成结论"
    ).strip()
    highlights = [
        f"当前状态：{case.get('state') or 'UNKNOWN'}",
        f"已纳入 {len(evidence_refs)} 条 Evidence，活跃假设 {len(active_hypotheses)} 个。",
    ]
    highlights.extend(
        f"假设：{item.get('statement') or item.get('description')}"
        for item in active_hypotheses[:3]
        if item.get("statement") or item.get("description")
    )
    highlights.extend(
        f"建议：{item.get('concrete_action') or item.get('title')}"
        for item in recommendations[:3]
        if item.get("concrete_action") or item.get("title")
    )
    summary = "\n\n".join([
        f"问题：{case.get('problem_description') or case.get('title') or case_id}",
        f"当前结论：{conclusion_text[:3000]}",
        "复盘要点：\n" + "\n".join(f"- {item}" for item in highlights),
    ])
    source_event_seq = max((int(item.get("case_event_seq") or 0) for item in events), default=0)
    now = now_utc()
    with new_session() as session:
        row = session.query(CaseMemoryModel).filter(
            CaseMemoryModel.case_id == case_id,
            CaseMemoryModel.tenant_id == tenant_id,
        ).first()
        if row is None:
            row = CaseMemoryModel(
                memory_id=f"memory-{secrets.token_hex(12)}",
                case_id=case_id,
                tenant_id=tenant_id,
                auto_capture=True,
                updated_at=now,
            )
            session.add(row)
        row.summary_text = summary[:MAX_CONTENT_CHARS]
        row.highlights = highlights[:12]
        row.evidence_refs = evidence_refs[:100]
        row.source_event_seq = source_event_seq
        row.generated_at = now
        row.updated_at = now
        session.commit()
        return row.to_dict()


def update_case_memory(case_id: str, tenant_id: str, *, auto_capture: bool) -> dict[str, Any] | None:
    with new_session() as session:
        row = session.query(CaseMemoryModel).filter(
            CaseMemoryModel.case_id == case_id,
            CaseMemoryModel.tenant_id == tenant_id,
        ).first()
        if row is None:
            return None
        row.auto_capture = bool(auto_capture)
        row.updated_at = now_utc()
        session.commit()
        return row.to_dict()


def promote_case_memory(repo: Any, case_id: str, tenant_id: str, created_by: str) -> dict[str, Any]:
    memory = get_case_memory(repo, case_id, tenant_id, refresh_auto=True)
    case = repo.get_incident_case(case_id, tenant_id) or {}
    document = create_knowledge_document(
        tenant_id=tenant_id,
        case_id=None,
        scope="GLOBAL",
        title=f"复盘：{case.get('title') or case_id}",
        filename=None,
        media_type="text/markdown",
        content_text=memory["summary_text"],
        created_by=created_by,
        kind="MEMORY",
    )
    with new_session() as session:
        row = session.query(CaseMemoryModel).filter(
            CaseMemoryModel.case_id == case_id,
            CaseMemoryModel.tenant_id == tenant_id,
        ).first()
        if row is not None:
            row.promoted_document_id = document["document_id"]
            row.updated_at = now_utc()
            session.commit()
    return document


def get_knowledge_document(document_id: str, tenant_id: str) -> dict[str, Any] | None:
    with new_session() as session:
        row = session.query(KnowledgeDocumentModel).filter(
            KnowledgeDocumentModel.document_id == document_id,
            KnowledgeDocumentModel.tenant_id == tenant_id,
        ).first()
        if row is None:
            return None
        value = _document_dict(session, row, include_content=True)
        chunks = session.query(KnowledgeChunkModel).filter(
            KnowledgeChunkModel.document_id == document_id,
        ).order_by(KnowledgeChunkModel.chunk_index).all()
        value["chunks"] = [chunk.to_dict() for chunk in chunks]
        return value


def get_knowledge_chunk(chunk_id: str, tenant_id: str, case_id: str | None = None) -> dict[str, Any] | None:
    with new_session() as session:
        pair = session.query(KnowledgeChunkModel, KnowledgeDocumentModel).join(
            KnowledgeDocumentModel,
            KnowledgeDocumentModel.document_id == KnowledgeChunkModel.document_id,
        ).filter(
            KnowledgeChunkModel.chunk_id == chunk_id,
            KnowledgeDocumentModel.tenant_id == tenant_id,
            KnowledgeDocumentModel.status == "ACTIVE",
        ).first()
        if pair is None:
            return None
        chunk, document = pair
        if document.scope != "GLOBAL" and document.case_id != case_id:
            return None
        return _chunk_result(1.0, chunk, document)


def chunk_text(value: str) -> list[tuple[int, int, str]]:
    """Create stable, overlapping character chunks while preferring paragraph boundaries."""
    text = str(value or "").strip()
    if not text:
        return []
    chunks: list[tuple[int, int, str]] = []
    start = 0
    length = len(text)
    while start < length:
        desired_end = min(length, start + CHUNK_TARGET_CHARS)
        end = desired_end
        if desired_end < length:
            boundary_start = min(length, start + CHUNK_TARGET_CHARS // 2)
            candidates = [
                text.rfind("\n\n", boundary_start, desired_end + 1),
                text.rfind("\n", boundary_start, desired_end + 1),
                text.rfind("。", boundary_start, desired_end + 1),
                text.rfind(". ", boundary_start, desired_end + 1),
            ]
            boundary = max(candidates)
            if boundary >= boundary_start:
                end = boundary + (2 if text[boundary:boundary + 2] in {"\n\n", ". "} else 1)
        content = text[start:end].strip()
        if content:
            content_start = start + len(text[start:end]) - len(text[start:end].lstrip())
            chunks.append((content_start, content_start + len(content), content))
        if end >= length:
            break
        next_start = max(start + 1, end - CHUNK_OVERLAP_CHARS)
        while next_start < end and not text[next_start].isspace():
            next_start += 1
        start = next_start if next_start < end else end
    return chunks


def _store_chunks(session: Any, document: KnowledgeDocumentModel, created_at: Any) -> None:
    chunks = chunk_text(document.content_text)
    vectors = embed_documents(content for _, _, content in chunks)
    for index, ((start, end, content), embedding) in enumerate(zip(chunks, vectors, strict=True)):
        terms = _tokens(f"{document.title} {content}")
        session.add(KnowledgeChunkModel(
            chunk_id=f"chunk-{secrets.token_hex(12)}",
            document_id=document.document_id,
            tenant_id=document.tenant_id,
            case_id=document.case_id,
            chunk_index=index,
            start_offset=start,
            end_offset=end,
            content_text=content,
            lexical_text=" ".join(terms),
            content_sha256=hashlib.sha256(content.encode("utf-8")).hexdigest(),
            term_frequencies=dict(Counter(terms)),
            embedding=embedding,
            created_at=created_at,
        ))


def _document_dict(session: Any, row: KnowledgeDocumentModel, *, include_content: bool = False) -> dict[str, Any]:
    value = row.to_dict(include_content=include_content)
    value["chunk_count"] = session.query(KnowledgeChunkModel).filter(
        KnowledgeChunkModel.document_id == row.document_id,
    ).count()
    value["index_status"] = "READY" if value["chunk_count"] else "EMPTY"
    value["embedding"] = embedding_status()
    return value


def _chunk_result(score: float, chunk: KnowledgeChunkModel, document: KnowledgeDocumentModel) -> dict[str, Any]:
    return {
        "knowledge_id": document.document_id,
        "document_id": document.document_id,
        "chunk_id": chunk.chunk_id,
        "chunk_index": chunk.chunk_index,
        "title": document.title,
        "content": chunk.content_text,
        "start_offset": chunk.start_offset,
        "end_offset": chunk.end_offset,
        "score": round(score, 6),
        "source": "operator_knowledge",
        "scope": document.scope,
        "case_id": document.case_id,
        "caveat": "用户知识只用于调查背景，不能替代当前运行时 Evidence。",
    }


def _limit(value: int) -> int:
    return max(1, min(int(value), 20))


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("DOCUMENT_ENCODING_UNSUPPORTED")


def _terms(value: str) -> set[str]:
    return set(_tokens(value))


def _tokens(value: str) -> list[str]:
    lowered = str(value or "").lower()
    terms = re.findall(r"[a-z0-9_.-]{2,}", lowered)
    for run in re.findall(r"[\u4e00-\u9fff]{2,}", lowered):
        terms.append(run)
        terms.extend(run[index:index + 2] for index in range(max(1, len(run) - 1)))
    return terms
